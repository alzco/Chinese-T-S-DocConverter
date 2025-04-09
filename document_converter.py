#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Document Converter Module
Handles conversion of text in various document formats (txt, md, docx)
using OpenCC for Chinese character conversion.
"""

import os
import io
import tempfile
import logging
import traceback
from docx import Document
import docx2txt
from opencc_converter import CustomOpenCC

# 设置日志记录
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')

class DocumentConverter:
    """
    A class for converting text in various document formats using OpenCC.
    Supports txt, md, and docx files.
    """
    
    def __init__(self, opencc_config='s2t', custom_dict=None):
        """
        Initialize the document converter.
        
        Args:
            opencc_config (str): The OpenCC configuration to use
            custom_dict (dict): Optional custom dictionary for conversion
        """
        self.converter = CustomOpenCC(opencc_config)
        
        # Add custom dictionary entries if provided
        if custom_dict:
            for source, target in custom_dict.items():
                self.converter.add_custom_mapping(source, target)
    
    def convert_text(self, text):
        """
        Convert plain text using OpenCC.
        
        Args:
            text (str): The text to convert
            
        Returns:
            str: The converted text
        """
        return self.converter.convert(text)
    
    def convert_txt_file(self, file_content):
        """
        Convert a text file.
        
        Args:
            file_content (bytes): The content of the text file
            
        Returns:
            str: The converted text
        """
        try:
            text = file_content.decode('utf-8')
            return self.convert_text(text)
        except UnicodeDecodeError:
            # Try with different encoding if UTF-8 fails
            try:
                text = file_content.decode('gbk')
                return self.convert_text(text)
            except:
                raise ValueError("无法解码文本文件，请确保文件编码为UTF-8或GBK")
    
    def convert_md_file(self, file_content):
        """
        Convert a Markdown file.
        
        Args:
            file_content (bytes): The content of the Markdown file
            
        Returns:
            str: The converted Markdown text
        """
        # Markdown files are treated the same as text files
        return self.convert_txt_file(file_content)
    
    def convert_docx_file(self, file_content):
        """
        Convert a DOCX file while preserving formatting.
        
        Args:
            file_content (bytes): The content of the DOCX file
            
        Returns:
            bytes: The content of the converted DOCX file
        """
        # Create a temporary file to store the input docx
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_in:
            temp_in.write(file_content)
            temp_in_path = temp_in.name
        
        try:
            logging.info("开始转换Word文档")
            # Load the document
            doc = Document(temp_in_path)
            
            # Convert text in paragraphs
            for paragraph in doc.paragraphs:
                self._convert_paragraph_runs(paragraph)
            
            # Convert text in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._convert_paragraph_runs(paragraph)
            
            # Convert text in footnotes - 使用更直接的方法处理脚注
            try:
                logging.info("开始处理脚注")
                
                # 检查文档中的脚注部分
                footnote_part = None
                for rel in doc.part.rels.values():
                    if rel.reltype == 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes':
                        footnote_part = rel.target_part
                        break
                
                if footnote_part:
                    logging.info(f"找到脚注部分: {footnote_part}")
                    # 直接修改脚注部分的XML
                    root = footnote_part.element.getroot()
                    nsmap = root.nsmap
                    w_namespace = nsmap.get('w') or '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                    
                    # 查找所有脚注
                    footnotes = root.findall(f'.//{w_namespace}footnote')
                    logging.info(f"找到 {len(footnotes)} 个脚注")
                    
                    for footnote in footnotes:
                        # 跳过分隔符脚注
                        footnote_id = footnote.get(f'{w_namespace}id')
                        if footnote_id in ['0', '-1']:
                            continue
                            
                        logging.info(f"处理脚注 ID: {footnote_id}")
                        
                        # 查找所有文本运行
                        text_elements = footnote.findall(f'.//{w_namespace}t')
                        for text_elem in text_elements:
                            if text_elem.text and text_elem.text.strip():
                                original_text = text_elem.text
                                converted_text = self.convert_text(original_text)
                                if converted_text != original_text:
                                    logging.info(f"脚注转换: '{original_text}' -> '{converted_text}'")
                                    text_elem.text = converted_text
                
                # 使用python-docx API作为备份方法
                if hasattr(doc, 'footnotes'):
                    footnotes_count = 0
                    for footnote in doc.footnotes._element.findall('.//' + footnote_part.element.nsmap['w'] + 'footnote'):
                        footnotes_count += 1
                    logging.info(f"python-docx API找到 {footnotes_count} 个脚注")
                    
                    for footnote in doc.footnotes.footnotes:
                        for paragraph in footnote.paragraphs:
                            for run in paragraph.runs:
                                if run.text:
                                    original_text = run.text
                                    converted_text = self.convert_text(original_text)
                                    if converted_text != original_text:
                                        logging.info(f"使用API转换脚注: '{original_text}' -> '{converted_text}'")
                                        run.text = converted_text
            except Exception as e:
                logging.error(f"处理脚注时出错: {e}", exc_info=True)
            
            # Convert text in endnotes - 使用更直接的方法处理尾注
            try:
                if hasattr(doc, '_element') and hasattr(doc._element, 'xpath'):
                    # 直接获取尾注内容的XML元素
                    endnote_elements = doc._element.xpath('.//w:endnote')
                    for endnote_element in endnote_elements:
                        # 获取尾注中的所有段落
                        for paragraph_element in endnote_element.xpath('.//w:p', namespaces=endnote_element.nsmap):
                            # 获取段落中的所有文本运行
                            for run_element in paragraph_element.xpath('.//w:t', namespaces=paragraph_element.nsmap):
                                if run_element.text:
                                    run_element.text = self.convert_text(run_element.text)
                
                # 使用标准方法作为备份
                if hasattr(doc, 'endnotes'):
                    for endnote in doc.endnotes:
                        for paragraph in endnote.paragraphs:
                            self._convert_paragraph_runs(paragraph)
            except Exception as e:
                print(f"处理尾注时出错: {e}")
            
            # Convert text in headers and footers
            for section in doc.sections:
                # Headers
                for header in [section.header]:
                    for paragraph in header.paragraphs:
                        self._convert_paragraph_runs(paragraph)
                    
                    for table in header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    self._convert_paragraph_runs(paragraph)
                
                # Footers
                for footer in [section.footer]:
                    for paragraph in footer.paragraphs:
                        self._convert_paragraph_runs(paragraph)
                    
                    for table in footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    self._convert_paragraph_runs(paragraph)
            
            # 输出调试信息
            logging.info("文档转换完成，准备保存")
            
            # Create a temporary file to store the output docx
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_out:
                temp_out_path = temp_out.name
            
            # Save the converted document
            doc.save(temp_out_path)
            logging.info(f"文档已保存到临时文件: {temp_out_path}")
            
            # Read the converted document
            with open(temp_out_path, 'rb') as f:
                converted_content = f.read()
            
            return converted_content
        
        except Exception as e:
            logging.error(f"转换Word文档时出错: {e}")
            logging.error(traceback.format_exc())
            raise
        finally:
            # Clean up temporary files
            if os.path.exists(temp_in_path):
                os.remove(temp_in_path)
            if 'temp_out_path' in locals() and os.path.exists(temp_out_path):
                os.remove(temp_out_path)
    
    def _convert_paragraph_runs(self, paragraph):
        """
        Convert text in paragraph runs while preserving formatting.
        
        Args:
            paragraph: The paragraph object to convert
        """
        if not paragraph.runs:
            # 如果没有runs但有文本，尝试转换整个段落文本
            if paragraph.text:
                # 保存原始文本
                original_text = paragraph.text
                # 转换文本
                converted_text = self.convert_text(original_text)
                # 如果转换成功且结果不同，则更新段落文本
                if converted_text != original_text:
                    paragraph.text = converted_text
            return
            
        # 处理段落中的每个文本运行
        for run in paragraph.runs:
            if run.text:
                # 保存原始文本
                original_text = run.text
                # 转换文本
                converted_text = self.convert_text(original_text)
                # 如果转换成功且结果不同，则更新文本运行
                if converted_text != original_text:
                    run.text = converted_text
    
    def convert_file(self, file_content, file_type):
        """
        Convert a file based on its type.
        
        Args:
            file_content (bytes): The content of the file
            file_type (str): The type of the file ('txt', 'md', or 'docx')
            
        Returns:
            Union[str, bytes]: The converted content (str for txt/md, bytes for docx)
        """
        if file_type == 'txt':
            return self.convert_txt_file(file_content)
        elif file_type == 'md':
            return self.convert_md_file(file_content)
        elif file_type == 'docx':
            return self.convert_docx_file(file_content)
        else:
            raise ValueError(f"不支持的文件类型: {file_type}")
